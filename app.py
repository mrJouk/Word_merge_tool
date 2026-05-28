from __future__ import annotations

import argparse
import logging
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable

import config as cfg


@dataclass(frozen=True)
class DocumentPair:
    name: str
    kyrgyz: Path
    russian: Path


@dataclass(frozen=True)
class WordBodyBlock:
    kind: str
    start: int
    end: int
    item: object


def parse_args() -> argparse.Namespace:
    """Read command-line options and apply defaults from config.py."""
    parser = argparse.ArgumentParser(
        description=(
            "Merge matched Kyrgyz and Russian Word documents into new two-column "
            "Word files while preserving as much original formatting as possible."
        )
    )
    parser.add_argument("--kyrgyz-dir", default=cfg.DEFAULT_KYRGYZ_DIR, help="Folder with Kyrgyz documents.")
    parser.add_argument("--russian-dir", default=cfg.DEFAULT_RUSSIAN_DIR, help="Folder with Russian documents.")
    parser.add_argument("--output-dir", default=cfg.DEFAULT_OUTPUT_DIR, help="Folder for merged .docx files.")
    parser.add_argument(
        "--backend",
        choices=cfg.BACKEND_CHOICES,
        default=cfg.DEFAULT_BACKEND,
        help=(
            "Merge backend. 'word' uses Microsoft Word automation and supports .doc/.docx. "
            "'docx' is a fallback for .docx only. 'auto' prefers Word and falls back to docx."
        ),
    )
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing merged files.")
    parser.add_argument(
        "--create-dummy-data",
        action="store_true",
        help="Create sample Kyrgyz and Russian .docx files, then exit.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Print detailed progress information.",
    )
    return parser.parse_args()


def setup_logging(verbose: bool) -> None:
    """Configure console logging for normal and verbose runs."""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(format=cfg.LOG_FORMAT, level=level)


def normalized_key(path: Path) -> str:
    """Use filename stem as the stable pairing key."""
    return path.stem.casefold()


def discover_documents(folder: Path) -> dict[str, list[Path]]:
    """Find supported Word files and group them by case-insensitive filename stem."""
    if not folder.exists():
        raise FileNotFoundError(f"Input folder does not exist: {folder}")
    if not folder.is_dir():
        raise NotADirectoryError(f"Input path is not a folder: {folder}")

    result: dict[str, list[Path]] = {}
    for path in sorted(folder.iterdir()):
        if path.is_file() and path.suffix.casefold() in cfg.SUPPORTED_EXTENSIONS:
            result.setdefault(normalized_key(path), []).append(path)
    return result


def choose_pair_member(files: list[Path], preferred_suffixes: Iterable[str]) -> Path:
    """Prefer a file whose extension matches the other language side."""
    preferred = {suffix.casefold() for suffix in preferred_suffixes}
    for path in files:
        if path.suffix.casefold() in preferred:
            return path
    return files[0]


def make_pairs(kyrgyz_dir: Path, russian_dir: Path) -> tuple[list[DocumentPair], list[str]]:
    """Pair Kyrgyz and Russian documents by matching filename stem."""
    kyrgyz_docs = discover_documents(kyrgyz_dir)
    russian_docs = discover_documents(russian_dir)
    warnings: list[str] = []
    pairs: list[DocumentPair] = []

    for key in sorted(set(kyrgyz_docs) & set(russian_docs)):
        kyrgyz_files = kyrgyz_docs[key]
        russian_files = russian_docs[key]
        russian_suffixes = [path.suffix for path in russian_files]
        kyrgyz = choose_pair_member(kyrgyz_files, russian_suffixes)
        russian = choose_pair_member(russian_files, [kyrgyz.suffix])

        if len(kyrgyz_files) > 1 or len(russian_files) > 1:
            warnings.append(
                "Multiple files matched stem "
                f"'{key}'. Using '{kyrgyz.name}' and '{russian.name}'."
            )

        pairs.append(DocumentPair(name=kyrgyz.stem, kyrgyz=kyrgyz, russian=russian))

    for key in sorted(set(kyrgyz_docs) - set(russian_docs)):
        warnings.append(f"No Russian match found for Kyrgyz file stem '{key}'.")
    for key in sorted(set(russian_docs) - set(kyrgyz_docs)):
        warnings.append(f"No Kyrgyz match found for Russian file stem '{key}'.")

    return pairs, warnings


def safe_output_name(name: str) -> str:
    """Build a filesystem-safe merged output filename."""
    cleaned = re.sub(cfg.SAFE_OUTPUT_NAME_PATTERN, "_", name).strip("._")
    return f"{cleaned or cfg.DEFAULT_SAFE_OUTPUT_STEM}{cfg.MERGED_OUTPUT_SUFFIX}"


class WordAutomationMerger:
    """High-fidelity merger that delegates legacy conversion and range copying to Word COM."""

    def __init__(self) -> None:
        self.word = None

    def __enter__(self) -> "WordAutomationMerger":
        try:
            import win32com.client  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "Microsoft Word backend requires pywin32. Install dependencies with "
                "`pip install -r requirements.txt`."
            ) from exc

        self.word = win32com.client.DispatchEx(cfg.WORD_APPLICATION_NAME)
        self.word.Visible = False
        self.word.DisplayAlerts = cfg.WORD_DISPLAY_ALERTS_NONE
        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        if self.word is not None:
            self.word.Quit()

    def merge(self, pair: DocumentPair, output_path: Path) -> None:
        """Create one merged document for a matched Kyrgyz/Russian pair."""
        if self.word is None:
            raise RuntimeError("WordAutomationMerger must be used as a context manager.")

        kyrgyz_doc = russian_doc = merged_doc = None
        try:
            kyrgyz_doc = self._open_source(pair.kyrgyz)
            russian_doc = self._open_source(pair.russian)
            merged_doc = self.word.Documents.Add()

            self._set_headers_and_footers(
                merged_doc,
                header_text=cfg.HEADER_FOOTER_SEPARATOR.join(
                    (
                        self._story_text(kyrgyz_doc, cfg.HEADER_STORY),
                        self._story_text(russian_doc, cfg.HEADER_STORY),
                    )
                ),
                footer_text=cfg.HEADER_FOOTER_SEPARATOR.join(
                    (
                        self._story_text(kyrgyz_doc, cfg.FOOTER_STORY),
                        self._story_text(russian_doc, cfg.FOOTER_STORY),
                    )
                ),
            )

            self._append_merged_body(kyrgyz_doc, russian_doc, merged_doc)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            merged_doc.SaveAs2(str(output_path.resolve()), FileFormat=cfg.WORD_FORMAT_XML_DOCUMENT)
            for doc in (merged_doc, kyrgyz_doc, russian_doc):
                doc.Close(SaveChanges=cfg.WORD_DO_NOT_SAVE_CHANGES)
            merged_doc = kyrgyz_doc = russian_doc = None
        finally:
            for doc in (merged_doc, kyrgyz_doc, russian_doc):
                if doc is not None:
                    doc.Close(SaveChanges=cfg.WORD_DO_NOT_SAVE_CHANGES)

    def _open_source(self, path: Path):
        return self.word.Documents.Open(
            FileName=str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            OpenAndRepair=True,
        )

    def _story_text(self, doc, story: str) -> str:
        """Collect visible header/footer text across all section variants."""
        pieces: list[str] = []

        for section in doc.Sections:
            collection = section.Headers if story == cfg.HEADER_STORY else section.Footers
            for index in cfg.WORD_HEADER_FOOTER_INDEXES:
                try:
                    text = cleanup_word_story_text(collection(index).Range.Text)
                except Exception:
                    continue
                if text:
                    pieces.append(text)

        return cfg.MULTI_SECTION_TEXT_SEPARATOR.join(dict.fromkeys(pieces))

    def _set_headers_and_footers(self, doc, header_text: str, footer_text: str) -> None:
        """Write merged header/footer text to all common Word header/footer variants."""
        for section in doc.Sections:
            for index in cfg.WORD_HEADER_FOOTER_INDEXES:
                section.Headers(index).Range.Text = header_text
                section.Footers(index).Range.Text = footer_text

    def _append_merged_body(self, kyrgyz_doc, russian_doc, merged_doc) -> None:
        """Copy non-table content with Word formatting and merge matching source tables."""
        kyrgyz_blocks = self._word_body_blocks(kyrgyz_doc)
        russian_blocks = self._word_body_blocks(russian_doc)
        kyrgyz_index = 0
        russian_index = 0

        while kyrgyz_index < len(kyrgyz_blocks) or russian_index < len(russian_blocks):
            kyrgyz_next = next_word_block(kyrgyz_blocks, kyrgyz_index)
            russian_next = next_word_block(russian_blocks, russian_index)

            if is_word_table_block(kyrgyz_next) and is_word_table_block(russian_next):
                self._append_combined_word_table(kyrgyz_next.item, russian_next.item, merged_doc)
                kyrgyz_index += 1
                russian_index += 1
                continue

            kyrgyz_text_blocks: list[WordBodyBlock] = []
            russian_text_blocks: list[WordBodyBlock] = []
            while kyrgyz_index < len(kyrgyz_blocks) and kyrgyz_blocks[kyrgyz_index].kind != "table":
                kyrgyz_text_blocks.append(kyrgyz_blocks[kyrgyz_index])
                kyrgyz_index += 1
            while russian_index < len(russian_blocks) and russian_blocks[russian_index].kind != "table":
                russian_text_blocks.append(russian_blocks[russian_index])
                russian_index += 1

            if kyrgyz_text_blocks or russian_text_blocks:
                self._append_two_column_word_ranges(
                    self._range_for_blocks(kyrgyz_doc, kyrgyz_text_blocks),
                    self._range_for_blocks(russian_doc, russian_text_blocks),
                    merged_doc,
                )
                continue

            if is_word_table_block(kyrgyz_next):
                self._append_two_column_word_ranges(kyrgyz_next.item.Range, None, merged_doc)
                kyrgyz_index += 1
            elif is_word_table_block(russian_next):
                self._append_two_column_word_ranges(None, russian_next.item.Range, merged_doc)
                russian_index += 1

    def _word_body_blocks(self, doc) -> list[WordBodyBlock]:
        """Return top-level paragraphs and tables in body order."""
        blocks: list[WordBodyBlock] = []
        for paragraph in doc.Paragraphs:
            try:
                if paragraph.Range.Information(cfg.WORD_INFO_WITHIN_TABLE):
                    continue
            except Exception:
                continue
            blocks.append(WordBodyBlock("text", paragraph.Range.Start, paragraph.Range.End, paragraph))

        for table in doc.Tables:
            blocks.append(WordBodyBlock("table", table.Range.Start, table.Range.End, table))

        return sorted(blocks, key=lambda block: (block.start, 0 if block.kind == "text" else 1))

    def _range_for_blocks(self, doc, blocks: list[WordBodyBlock]):
        """Create one Word range spanning adjacent non-table blocks."""
        if not blocks:
            return None
        start = min(block.start for block in blocks)
        end = max(block.end for block in blocks)
        return doc.Range(start, end)

    def _append_two_column_word_ranges(self, kyrgyz_range, russian_range, merged_doc) -> None:
        """Append formatted source ranges in the standard Kyrgyz/Russian two-column table."""
        table = self._add_word_table_at_end(merged_doc, cfg.MERGED_TABLE_ROWS, cfg.MERGED_TABLE_COLUMNS)
        table.Borders.Enable = True
        self._fix_two_column_word_table_widths(merged_doc, table)
        self._copy_range_to_cell(kyrgyz_range, table.Cell(cfg.WORD_TABLE_ROW, cfg.WORD_KYRGYZ_COLUMN))
        self._copy_range_to_cell(russian_range, table.Cell(cfg.WORD_TABLE_ROW, cfg.WORD_RUSSIAN_COLUMN))
        self._fix_two_column_word_table_widths(merged_doc, table)
        self._append_paragraph_after_word_table(merged_doc, table)

    def _copy_range_to_cell(self, source_range, cell) -> None:
        """Copy a Word COM range as formatted text into a table cell."""
        if source_range is None:
            return
        target_range = cell.Range
        if target_range.End > target_range.Start:
            target_range.End = target_range.End - 1
        target_range.FormattedText = source_range.FormattedText

    def _append_combined_word_table(self, kyrgyz_table, russian_table, merged_doc) -> None:
        """Create one Word table from matching source tables using combined cell text."""
        row_count = max(word_table_row_count(kyrgyz_table), word_table_row_count(russian_table), 1)
        column_count = max(word_table_column_count(kyrgyz_table), word_table_column_count(russian_table), 1)
        table = self._add_word_table_at_end(merged_doc, row_count, column_count)
        table.Borders.Enable = True
        table.AllowAutoFit = True

        for row_index in range(1, row_count + 1):
            for column_index in range(1, column_count + 1):
                kyrgyz_text = word_table_cell_text(kyrgyz_table, row_index, column_index)
                russian_text = word_table_cell_text(russian_table, row_index, column_index)
                cell = table.Cell(row_index, column_index)
                cell.Range.Text = combine_cell_text(kyrgyz_text, russian_text)
                if row_index == cfg.WORD_TABLE_ROW:
                    cell.Range.Bold = True
        self._append_paragraph_after_word_table(merged_doc, table)

    def _add_word_table_at_end(self, doc, rows: int, columns: int):
        """Insert a Word table at the end of the destination document."""
        if doc.Tables.Count > 0:
            end_range = doc.Range(max(doc.Content.End - 1, doc.Content.Start), max(doc.Content.End - 1, doc.Content.Start))
            end_range.InsertParagraphAfter()
        insert_at = max(doc.Content.End - 1, doc.Content.Start)
        target_range = doc.Range(insert_at, insert_at)
        return doc.Tables.Add(target_range, rows, columns)

    def _fix_two_column_word_table_widths(self, doc, table) -> None:
        """Force comparison tables to full width with equal Kyrgyz/Russian columns."""
        try:
            table.AllowAutoFit = False
            table.AutoFitBehavior(cfg.WORD_AUTOFIT_FIXED)
        except Exception as exc:
            logging.debug("Could not disable Word table autofit: %s", exc)

        try:
            usable_width = doc.PageSetup.PageWidth - doc.PageSetup.LeftMargin - doc.PageSetup.RightMargin
            column_width = usable_width * cfg.TWO_COLUMN_WIDTH_RATIO
            table.PreferredWidthType = cfg.WORD_PREFERRED_WIDTH_POINTS
            table.PreferredWidth = usable_width
            table.Columns(cfg.WORD_KYRGYZ_COLUMN).Width = column_width
            table.Columns(cfg.WORD_RUSSIAN_COLUMN).Width = column_width
        except Exception as exc:
            logging.debug("Could not set equal Word table column widths: %s", exc)

        try:
            table.Columns.DistributeWidth()
        except Exception as exc:
            logging.debug("Could not distribute Word table columns: %s", exc)

    def _append_paragraph_after_word_table(self, doc, table) -> None:
        """Create a real paragraph after a table so the next table is not nested."""
        after_table = doc.Range(table.Range.End, table.Range.End)
        after_table.InsertParagraphAfter()


def cleanup_word_story_text(text: str) -> str:
    """Remove Word story control characters and normalize header/footer whitespace."""
    text = text.replace("\x07", "").replace("\r", "\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return " ".join(lines)


def next_word_block(blocks: list[WordBodyBlock], index: int) -> WordBodyBlock | None:
    """Return the next Word COM body block or None."""
    if index >= len(blocks):
        return None
    return blocks[index]


def is_word_table_block(block: WordBodyBlock | None) -> bool:
    """Check whether a Word COM block is a table."""
    return block is not None and block.kind == "table"


def word_table_row_count(table) -> int:
    """Return a Word COM table row count, tolerating irregular tables."""
    try:
        return table.Rows.Count
    except Exception:
        return 0


def word_table_column_count(table) -> int:
    """Return a Word COM table column count, tolerating irregular tables."""
    try:
        return table.Columns.Count
    except Exception:
        counts: list[int] = []
        for row_index in range(1, word_table_row_count(table) + 1):
            try:
                counts.append(table.Rows(row_index).Cells.Count)
            except Exception:
                continue
        return max(counts, default=0)


def word_table_cell_text(table, row_index: int, column_index: int) -> str:
    """Extract normalized visible text from a Word COM table cell."""
    try:
        text = table.Cell(row_index, column_index).Range.Text
    except Exception:
        return ""
    return cleanup_word_story_text(text)


class DocxFallbackMerger:
    """Fallback merger for .docx inputs when Microsoft Word automation is unavailable."""

    def merge(self, pair: DocumentPair, output_path: Path) -> None:
        if (
            pair.kyrgyz.suffix.casefold() == cfg.LEGACY_DOC_EXTENSION
            or pair.russian.suffix.casefold() == cfg.LEGACY_DOC_EXTENSION
        ):
            raise RuntimeError(".doc files require the Microsoft Word backend.")

        DocxStructuralMerger().merge_paths(pair.kyrgyz, pair.russian, output_path)


class DocxStructuralMerger:
    """Merge prepared .docx files while pairing corresponding source tables."""

    def __init__(self) -> None:
        try:
            from docx import Document  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
            from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "The .docx merge step requires python-docx. Install dependencies with "
                "`pip install -r requirements.txt`."
            ) from exc

        self.Document = Document
        self.qn = qn
        self.RT = RT

    def merge_paths(self, kyrgyz_path: Path, russian_path: Path, output_path: Path) -> None:
        """Create the final document from two prepared .docx sources."""
        kyrgyz_doc = self.Document(kyrgyz_path)
        russian_doc = self.Document(russian_path)
        merged_doc = self.Document()

        self._set_combined_headers_and_footers(kyrgyz_doc, russian_doc, merged_doc)
        copy_missing_styles(kyrgyz_doc, merged_doc, self.qn)
        copy_missing_styles(russian_doc, merged_doc, self.qn)
        copy_missing_numbering(kyrgyz_doc, merged_doc, self.qn)
        copy_missing_numbering(russian_doc, merged_doc, self.qn)
        clear_document_body(merged_doc, self.qn)
        self._append_merged_body(kyrgyz_doc, russian_doc, merged_doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        merged_doc.save(output_path)

    def _set_combined_headers_and_footers(self, kyrgyz_doc, russian_doc, merged_doc) -> None:
        """Combine header and footer text from both source documents."""
        section = merged_doc.sections[0]
        section.header.paragraphs[0].text = cfg.HEADER_FOOTER_SEPARATOR.join(
            (
                docx_story_text(kyrgyz_doc, cfg.HEADER_STORY, self.qn),
                docx_story_text(russian_doc, cfg.HEADER_STORY, self.qn),
            )
        )
        section.footer.paragraphs[0].text = cfg.HEADER_FOOTER_SEPARATOR.join(
            (
                docx_story_text(kyrgyz_doc, cfg.FOOTER_STORY, self.qn),
                docx_story_text(russian_doc, cfg.FOOTER_STORY, self.qn),
            )
        )

    def _append_merged_body(self, kyrgyz_doc, russian_doc, merged_doc) -> None:
        """Append non-table content in two columns and matching tables as one table."""
        kyrgyz_blocks = body_blocks(kyrgyz_doc, self.qn)
        russian_blocks = body_blocks(russian_doc, self.qn)
        kyrgyz_index = 0
        russian_index = 0

        while kyrgyz_index < len(kyrgyz_blocks) or russian_index < len(russian_blocks):
            kyrgyz_next = next_block(kyrgyz_blocks, kyrgyz_index)
            russian_next = next_block(russian_blocks, russian_index)

            if is_table_block(kyrgyz_next, self.qn) and is_table_block(russian_next, self.qn):
                self._append_merged_table(kyrgyz_next, russian_next, merged_doc)
                kyrgyz_index += 1
                russian_index += 1
                continue

            kyrgyz_nodes: list = []
            russian_nodes: list = []

            while kyrgyz_index < len(kyrgyz_blocks) and not is_table_block(kyrgyz_blocks[kyrgyz_index], self.qn):
                kyrgyz_nodes.append(kyrgyz_blocks[kyrgyz_index])
                kyrgyz_index += 1
            while russian_index < len(russian_blocks) and not is_table_block(russian_blocks[russian_index], self.qn):
                russian_nodes.append(russian_blocks[russian_index])
                russian_index += 1

            if kyrgyz_nodes or russian_nodes:
                self._append_two_column_content(kyrgyz_nodes, russian_nodes, kyrgyz_doc, russian_doc, merged_doc)
                continue

            if is_table_block(kyrgyz_next, self.qn):
                self._append_two_column_content([kyrgyz_next], [], kyrgyz_doc, russian_doc, merged_doc)
                kyrgyz_index += 1
            elif is_table_block(russian_next, self.qn):
                self._append_two_column_content([], [russian_next], kyrgyz_doc, russian_doc, merged_doc)
                russian_index += 1

    def _append_two_column_content(self, kyrgyz_nodes, russian_nodes, kyrgyz_doc, russian_doc, merged_doc) -> None:
        """Append unmatched text/table blocks in the standard Kyrgyz/Russian two-column layout."""
        table = merged_doc.add_table(rows=cfg.MERGED_TABLE_ROWS, cols=cfg.MERGED_TABLE_COLUMNS)
        table.style = cfg.TABLE_GRID_STYLE
        self._fix_two_column_docx_table_widths(merged_doc, table)
        copy_nodes_to_cell(kyrgyz_nodes, kyrgyz_doc.part, merged_doc.part, table.cell(0, 0), self.qn, self.RT)
        copy_nodes_to_cell(russian_nodes, russian_doc.part, merged_doc.part, table.cell(0, 1), self.qn, self.RT)
        self._fix_two_column_docx_table_widths(merged_doc, table)

    def _fix_two_column_docx_table_widths(self, doc, table) -> None:
        """Force fallback comparison tables to fixed equal columns."""
        section = doc.sections[0]
        column_width = int((section.page_width - section.left_margin - section.right_margin) * cfg.TWO_COLUMN_WIDTH_RATIO)
        table.autofit = False
        for column in table.columns:
            column.width = column_width
        for row in table.rows:
            for cell in row.cells:
                cell.width = column_width

    def _append_merged_table(self, kyrgyz_table, russian_table, merged_doc) -> None:
        """Create one table from matching source tables, combining cell text with '/'."""
        row_count = max(table_row_count(kyrgyz_table, self.qn), table_row_count(russian_table, self.qn))
        column_count = max(table_column_count(kyrgyz_table, self.qn), table_column_count(russian_table, self.qn))
        table = merged_doc.add_table(rows=max(row_count, 1), cols=max(column_count, 1))
        table.style = cfg.TABLE_GRID_STYLE

        for row_index in range(max(row_count, 1)):
            for column_index in range(max(column_count, 1)):
                kyrgyz_text = table_cell_text(kyrgyz_table, row_index, column_index, self.qn)
                russian_text = table_cell_text(russian_table, row_index, column_index, self.qn)
                table.cell(row_index, column_index).text = combine_cell_text(kyrgyz_text, russian_text)
                if row_index == cfg.DOCX_TABLE_ROW:
                    for paragraph in table.cell(row_index, column_index).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def docx_story_text(doc, story: str, qn) -> str:
    """Extract plain header/footer text from a python-docx document."""
    pieces: list[str] = []
    for section in doc.sections:
        source = section.header if story == cfg.HEADER_STORY else section.footer
        words = [node.text for node in source._element.iter(qn("w:t")) if node.text]
        text = " ".join(" ".join(words).split())
        if text:
            pieces.append(text)
    return cfg.MULTI_SECTION_TEXT_SEPARATOR.join(dict.fromkeys(pieces))


def body_blocks(doc, qn) -> list:
    """Return body XML blocks except section properties."""
    return [child for child in doc.element.body.iterchildren() if child.tag != qn("w:sectPr")]


def next_block(blocks: list, index: int):
    """Return the next body block or None when the sequence is exhausted."""
    if index >= len(blocks):
        return None
    return blocks[index]


def is_table_block(block, qn) -> bool:
    """Check whether a body block is a Word table XML element."""
    return block is not None and block.tag == qn("w:tbl")


def clear_document_body(doc, qn) -> None:
    """Remove default body content while preserving document section properties."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def copy_nodes_to_cell(nodes: list, source_part, target_part, cell, qn, RT) -> None:
    """Clone selected XML body nodes into a table cell."""
    clear_cell(cell)
    copied_any = False
    for node in nodes:
        cloned = deepcopy(node)
        remap_relationships(source_part, target_part, cloned, qn, RT)
        cell._tc.append(cloned)
        copied_any = True

    if not copied_any:
        cell.add_paragraph("")


def table_row_count(table_node, qn) -> int:
    """Count rows in a table XML node."""
    return len(table_node.findall(qn("w:tr")))


def table_column_count(table_node, qn) -> int:
    """Count the maximum number of cells found in any table row."""
    counts = [len(row.findall(qn("w:tc"))) for row in table_node.findall(qn("w:tr"))]
    return max(counts, default=0)


def table_cell_text(table_node, row_index: int, column_index: int, qn) -> str:
    """Extract normalized text from one table cell by row and column index."""
    rows = table_node.findall(qn("w:tr"))
    if row_index >= len(rows):
        return ""
    cells = rows[row_index].findall(qn("w:tc"))
    if column_index >= len(cells):
        return ""

    text_parts = [node.text for node in cells[column_index].iter(qn("w:t")) if node.text]
    return " ".join(" ".join(text_parts).split())


def combine_cell_text(kyrgyz_text: str, russian_text: str) -> str:
    """Combine matching table cell text from the two source documents."""
    if kyrgyz_text and russian_text:
        return cfg.HEADER_FOOTER_SEPARATOR.join((kyrgyz_text, russian_text))
    return kyrgyz_text or russian_text


def copy_missing_styles(source_doc, target_doc, qn) -> None:
    """Copy style definitions referenced by cloned .docx body XML."""
    target_styles = target_doc.styles.element
    existing = {
        style.get(qn("w:styleId"))
        for style in target_styles.findall(qn("w:style"))
        if style.get(qn("w:styleId"))
    }
    for style in source_doc.styles.element.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        if style_id and style_id not in existing:
            target_styles.append(deepcopy(style))
            existing.add(style_id)


def copy_missing_numbering(source_doc, target_doc, qn) -> None:
    """Copy numbering definitions so lists keep their formatting in cloned XML."""
    try:
        source_numbering = source_doc.part.numbering_part.element
        target_numbering = target_doc.part.numbering_part.element
    except Exception:
        return

    existing_abstract = {
        node.get(qn("w:abstractNumId"))
        for node in target_numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId"))
    }
    existing_nums = {
        node.get(qn("w:numId"))
        for node in target_numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    }

    for node in source_numbering.findall(qn("w:abstractNum")):
        node_id = node.get(qn("w:abstractNumId"))
        if node_id and node_id not in existing_abstract:
            target_numbering.append(deepcopy(node))
            existing_abstract.add(node_id)

    for node in source_numbering.findall(qn("w:num")):
        node_id = node.get(qn("w:numId"))
        if node_id and node_id not in existing_nums:
            target_numbering.append(deepcopy(node))
            existing_nums.add(node_id)


def clear_cell(cell) -> None:
    """Remove the default empty paragraph before inserting cloned content."""
    for child in list(cell._tc):
        cell._tc.remove(child)


def remap_relationships(source_part, target_part, element, qn, RT) -> None:
    """Recreate image and external relationships after XML is copied between parts."""
    relationship_attributes = (qn("r:id"), qn("r:embed"), qn("r:link"))
    for node in element.iter():
        for attr in relationship_attributes:
            old_rid = node.get(attr)
            if not old_rid or old_rid not in source_part.rels:
                continue

            rel = source_part.rels[old_rid]
            if rel.is_external:
                new_rid = target_part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            elif rel.reltype == RT.IMAGE:
                new_rid, _ = target_part.get_or_add_image(BytesIO(rel.target_part.blob))
            else:
                new_rid = target_part.relate_to(rel.target_part, rel.reltype)
            node.set(attr, new_rid)


def create_dummy_data(kyrgyz_dir: Path, russian_dir: Path) -> None:
    """Create representative .docx samples and optional legacy .doc samples."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import RGBColor  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "Dummy data generation requires python-docx. Install dependencies with "
            "`pip install -r requirements.txt`."
        ) from exc

    kyrgyz_dir.mkdir(parents=True, exist_ok=True)
    russian_dir.mkdir(parents=True, exist_ok=True)
    build_dummy_doc(
        Document(),
        kyrgyz_dir / cfg.DUMMY_KYRGYZ_DOCX,
        header=cfg.DUMMY_KYRGYZ["header"],
        footer=cfg.DUMMY_KYRGYZ["footer"],
        title=cfg.DUMMY_KYRGYZ["title"],
        lead=cfg.DUMMY_KYRGYZ["lead"],
        table_values=cfg.DUMMY_KYRGYZ["table_values"],
        accent=RGBColor(*cfg.DUMMY_KYRGYZ["accent_rgb"]),
    )
    build_dummy_doc(
        Document(),
        russian_dir / cfg.DUMMY_RUSSIAN_DOCX,
        header=cfg.DUMMY_RUSSIAN["header"],
        footer=cfg.DUMMY_RUSSIAN["footer"],
        title=cfg.DUMMY_RUSSIAN["title"],
        lead=cfg.DUMMY_RUSSIAN["lead"],
        table_values=cfg.DUMMY_RUSSIAN["table_values"],
        accent=RGBColor(*cfg.DUMMY_RUSSIAN["accent_rgb"]),
    )
    create_legacy_doc_fixtures(kyrgyz_dir, russian_dir)


def create_legacy_doc_fixtures(kyrgyz_dir: Path, russian_dir: Path) -> None:
    """Use Word COM to create legacy .doc fixtures from the generated .docx files."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        logging.warning("pywin32 is unavailable; skipping legacy .doc dummy files.")
        return

    word = None
    try:
        word = win32com.client.DispatchEx(cfg.WORD_APPLICATION_NAME)
        word.Visible = False
        word.DisplayAlerts = cfg.WORD_DISPLAY_ALERTS_NONE
        for source, target in (
            (kyrgyz_dir / cfg.DUMMY_KYRGYZ_DOCX, kyrgyz_dir / cfg.DUMMY_LEGACY_DOC),
            (russian_dir / cfg.DUMMY_RUSSIAN_DOCX, russian_dir / cfg.DUMMY_LEGACY_DOC),
        ):
            doc = word.Documents.Open(str(source.resolve()), ReadOnly=True, AddToRecentFiles=False)
            try:
                doc.SaveAs2(str(target.resolve()), FileFormat=cfg.WORD_FORMAT_DOC)
            finally:
                doc.Close(SaveChanges=cfg.WORD_DO_NOT_SAVE_CHANGES)
    except Exception as exc:
        logging.warning("Microsoft Word is unavailable; skipping legacy .doc dummy files. Reason: %s", exc)
    finally:
        if word is not None:
            word.Quit()


def build_dummy_doc(doc, path: Path, header: str, footer: str, title: str, lead: str, table_values, accent) -> None:
    """Build one formatted sample document with body text, a table, header, and footer."""
    section = doc.sections[0]
    section.header.paragraphs[0].text = header
    section.footer.paragraphs[0].text = footer

    doc.add_heading(title, level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run(lead + " ")
    bold = paragraph.add_run("Bold")
    bold.bold = True
    paragraph.add_run(" / ")
    italic = paragraph.add_run("Italic")
    italic.italic = True
    paragraph.add_run(" / ")
    colored = paragraph.add_run("Color")
    colored.font.color.rgb = accent

    table = doc.add_table(rows=len(table_values), cols=2)
    table.style = cfg.TABLE_GRID_STYLE
    for row_index, row_values in enumerate(table_values):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            run = cell.paragraphs[0].add_run(value)
            if row_index == 0:
                run.bold = True

    doc.add_paragraph("Final paragraph after the table.")
    doc.save(path)


def merge_all(
    pairs: list[DocumentPair],
    output_dir: Path,
    backend: str,
    overwrite: bool,
) -> int:
    """Select the requested backend and merge every discovered pair."""
    if not pairs:
        logging.warning("No matching document pairs found.")
        return 0

    output_dir.mkdir(parents=True, exist_ok=True)

    if backend == cfg.BACKEND_WORD:
        with WordAutomationMerger() as merger:
            return run_merges(pairs, output_dir, merger, overwrite)

    if backend == cfg.BACKEND_DOCX:
        return run_merges(pairs, output_dir, DocxFallbackMerger(), overwrite)

    try:
        with WordAutomationMerger() as merger:
            return run_merges(pairs, output_dir, merger, overwrite)
    except Exception as exc:
        if any(path.suffix.casefold() == cfg.LEGACY_DOC_EXTENSION for pair in pairs for path in (pair.kyrgyz, pair.russian)):
            raise RuntimeError(
                "Microsoft Word backend is required because at least one matched file is .doc. "
                f"Word backend failed with: {exc}"
            ) from exc
        logging.warning("Word backend unavailable; using .docx fallback. Reason: %s", exc)
        return run_merges(pairs, output_dir, DocxFallbackMerger(), overwrite)


def run_merges(pairs: list[DocumentPair], output_dir: Path, merger, overwrite: bool) -> int:
    """Run a merger implementation for each pair and count created files."""
    merged_count = 0
    for pair in pairs:
        output_path = output_dir / safe_output_name(pair.name)
        if output_path.exists() and not overwrite:
            logging.info("Skipping existing output: %s", output_path)
            continue

        logging.info("Merging %s + %s -> %s", pair.kyrgyz.name, pair.russian.name, output_path.name)
        merger.merge(pair, output_path)
        merged_count += 1

    return merged_count


def main() -> int:
    """Application entry point."""
    args = parse_args()
    setup_logging(args.verbose)

    kyrgyz_dir = Path(args.kyrgyz_dir)
    russian_dir = Path(args.russian_dir)
    output_dir = Path(args.output_dir)

    try:
        if args.create_dummy_data:
            create_dummy_data(kyrgyz_dir, russian_dir)
            logging.info("Dummy data created in %s and %s.", kyrgyz_dir, russian_dir)
            return 0

        pairs, warnings = make_pairs(kyrgyz_dir, russian_dir)
        for warning in warnings:
            logging.warning(warning)
        logging.info("Found %s matching document pair(s).", len(pairs))

        merged_count = merge_all(pairs, output_dir, args.backend, args.overwrite)
        logging.info("Created %s merged document(s).", merged_count)
        return 0
    except Exception as exc:
        logging.error("%s", exc)
        if args.verbose:
            raise
        return 1


if __name__ == "__main__":
    sys.exit(main())
